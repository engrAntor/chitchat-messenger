import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

doc = docx.Document()

# Page Margins (1 inch all sides)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

PRIMARY_COLOR = RGBColor(79, 70, 229)    # Indigo 600
DARK_COLOR = RGBColor(30, 41, 59)        # Slate 800
MUTED_COLOR = RGBColor(100, 116, 139)    # Slate 500
BORDER_COLOR = "E2E8F0"

def add_header(title, level=1):
    h = doc.add_heading(title, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    for run in h.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.color.rgb = PRIMARY_COLOR
            run.font.size = Pt(18)
            run.bold = True
        elif level == 2:
            run.font.color.rgb = DARK_COLOR
            run.font.size = Pt(14)
            run.bold = True
        elif level == 3:
            run.font.color.rgb = DARK_COLOR
            run.font.size = Pt(12)
            run.bold = True
    return h

# Title
title_p = doc.add_paragraph()
title_p.paragraph_format.space_after = Pt(4)
r_title = title_p.add_run("Frontend Developer Task Submission")
r_title.font.name = 'Calibri'
r_title.font.size = Pt(24)
r_title.bold = True
r_title.font.color.rgb = PRIMARY_COLOR

# Subtitle
sub_p = doc.add_paragraph()
sub_p.paragraph_format.space_after = Pt(18)
r_sub = sub_p.add_run("Project: Real-Time Chat Application (ChitChat Messenger) | Candidate: Antor Chandra Das")
r_sub.font.name = 'Calibri'
r_sub.font.size = Pt(11)
r_sub.font.color.rgb = MUTED_COLOR

# Section 1: Submission Email
add_header("1. Official Submission Email", level=1)

email_box = doc.add_table(rows=1, cols=1)
email_box.alignment = WD_TABLE_ALIGNMENT.CENTER
email_box.autofit = False
cell = email_box.cell(0, 0)
cell.width = Inches(6.5)

# Style email cell
shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8FAFC"/>')
cell._tc.get_or_add_tcPr().append(shading)
borders = parse_xml(
    f'<w:tcBorders {nsdecls("w")}>'
    f'<w:top w:val="single" w:sz="8" w:space="0" w:color="{BORDER_COLOR}"/>'
    f'<w:left w:val="single" w:sz="24" w:space="0" w:color="4F46E5"/>'
    f'<w:bottom w:val="single" w:sz="8" w:space="0" w:color="{BORDER_COLOR}"/>'
    f'<w:right w:val="single" w:sz="8" w:space="0" w:color="{BORDER_COLOR}"/>'
    f'</w:tcBorders>'
)
cell._tc.get_or_add_tcPr().append(borders)

email_text = (
    "Subject: Submission: Frontend Developer Task – Real-Time Chat Application – Antor Chandra Das\n\n"
    "Dear Hiring Team,\n\n"
    "I hope this email finds you well.\n\n"
    "I have completed the Frontend Developer take-home assignment and am pleased to submit my work for review. Below are the project links and an overview of the implementation deliverables:\n\n"
    "🔗 PROJECT LINKS\n"
    "• Live Demo (Netlify): https://engr-antor-chat-app.netlify.app\n"
    "• GitHub Repository: https://github.com/engrAntor/chitchat-messenger\n"
    "• API Documentation: Included in the repository under API_DOCUMENTATION.md\n\n"
    "📌 SUMMARY OF DELIVERABLES\n\n"
    "Part 1: API Documentation (API_DOCUMENTATION.md)\n"
    "• Comprehensive REST endpoints specification (/auth, /users, /conversations, /messages).\n"
    "• Request/response schemas, JWT bearer authorization flows, and error handling matrices.\n"
    "• Real-time Socket.IO event documentation (setup, room joins, bidirectional messaging).\n\n"
    "Part 2: Real-Time Web Application (Next.js & TypeScript)\n"
    "• Sub-second bidirectional messaging for 1-on-1 and group chats.\n"
    "• Optimistic UI delivery with UUID reconciliation and error-retry flows.\n"
    "• Group chat management: creation, user search, admin role promotion, and member removal.\n"
    "• Responsive UI with Dark/Light modes, smart auto-scroll, and unread badges.\n\n"
    "Part 3: Architecture & Documentation (README.md)\n"
    "• Deep-dive technical rationale for TypeScript, Next.js 16 (App Router), Zustand, and Tailwind CSS.\n"
    "• Step-by-step local development, Vercel, and Netlify deployment instructions.\n\n"
    "Thank you for the opportunity to work on this assignment. I look forward to your feedback!\n\n"
    "Best regards,\n"
    "Antor Chandra Das\n"
    "Frontend Developer\n"
    "GitHub: https://github.com/engrAntor"
)

cell_p = cell.paragraphs[0]
cell_p.paragraph_format.space_before = Pt(6)
cell_p.paragraph_format.space_after = Pt(6)
r = cell_p.add_run(email_text)
r.font.name = 'Consolas'
r.font.size = Pt(9.5)
r.font.color.rgb = DARK_COLOR

# Section 2: Key Links
add_header("2. Project References & Links", level=1)
links_table = doc.add_table(rows=5, cols=2)
links_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Resource", "URL / Location"]
row_data = [
    ["Live Production Deployment", "https://engr-antor-chat-app.netlify.app"],
    ["GitHub Source Code", "https://github.com/engrAntor/chitchat-messenger"],
    ["Backend API Base URL", "https://frontend-task-chatapp.onrender.com"],
    ["Swagger / OpenAPI Docs", "https://frontend-task-chatapp.onrender.com/docs/"],
]

for col_idx, text in enumerate(headers):
    c = links_table.cell(0, col_idx)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EEF2FF"/>')
    c._tc.get_or_add_tcPr().append(shading)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = PRIMARY_COLOR

for row_idx, data in enumerate(row_data):
    for col_idx, text in enumerate(data):
        c = links_table.cell(row_idx + 1, col_idx)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARK_COLOR

# Section 3: Technical Highlights
add_header("3. Architecture & Technical Highlights", level=1)

p = doc.add_paragraph()
r = p.add_run("Key Architectural Decisions:")
r.bold = True

bullets = [
    ("TypeScript Adoption: ", "Enforced 100% strict type safety across REST contracts, Socket.IO payload definitions, and UI state slices, completely eliminating runtime null-pointer exceptions."),
    ("Next.js 16 (App Router): ", "Utilized modern React Server Components where beneficial, coupled with optimized client-side hydration for dynamic real-time communication."),
    ("Zustand State Architecture: ", "Engineered isolated store slices for authentication and chat state with optimistic update reconciliation, unread badge counters, and cursor-based pagination."),
    ("Resilient Socket.IO Streams: ", "Implemented auto-reconnecting WebSocket transports with HTTP long-polling fallback, multi-event listeners, and automated room subscriptions upon login and navigation.")
]

for b_title, b_desc in bullets:
    bp = doc.add_paragraph(style='List Bullet')
    bp.paragraph_format.space_after = Pt(3)
    r1 = bp.add_run(b_title)
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.color.rgb = DARK_COLOR
    r2 = bp.add_run(b_desc)
    r2.font.name = 'Calibri'
    r2.font.color.rgb = DARK_COLOR

output_path = r"e:\chat-app\PROJECT_SUBMISSION.docx"
doc.save(output_path)
print("Saved document to:", output_path)
